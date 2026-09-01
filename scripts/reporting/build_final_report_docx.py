#!/usr/bin/env python3
"""Build the Thai project-summary report as a Word document on the department form.

Reads three inputs and writes one .docx.  Nothing upstream is modified.

  1. docs/source/report_power_system_project_final_th_v2.tex
     The report body.  Every sentence, number, table cell and caption comes from
     here; this script does not author content.

  2. docs/source/report_power_system_project_final_th_v2.aux
     The reference-number map LaTeX already computed (figures 1-7, tables 1-10,
     equations 1-34 plus the tagged (M), citations [1]-[16]).  Word has no
     equivalent of amsmath auto-numbering, so those numbers are lifted from the
     .aux verbatim and the Word file therefore agrees with the PDF digit for
     digit.  IF THE .TEX CHANGES SO THAT NUMBERING SHIFTS, RUN xelatex FIRST so
     the .aux is current, then re-run this script.

  3. docs/text/แบบฟอร์ม เอกสารสรุปวิชาโครงงาน.docx
     The department form, used as the base of the output so that its styles,
     fonts, footer and page setup are the ones that ship in the deliverable.
     The form as distributed carries a [trash]/0000.dat part that is absent from
     [Content_Types].xml, which is why Word reports the file as damaged.  Saving
     through python-docx drops that part, so the output opens clean.

PROVENANCE AND NUMERICAL INTEGRITY
  - No numeric value is recomputed, rounded, reformatted or re-derived here.
    Table cells and inline quantities are copied out of the .tex as written.
  - The two TikZ figures are rendered from the .tex source itself via xelatex,
    not redrawn, so they carry the same \\NewRun* macro values as the PDF.
  - Raster figures are the PNGs the MATLAB generator produced; they are scaled
    to the form's narrower text column but never regenerated or edited.
  - Every equation must convert.  A conversion failure aborts the build rather
    than silently dropping the equation.
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

import latex2mathml.converter
import win32com.client

# ---------------------------------------------------------------------------
# paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "source"
TEX = SRC / "report_power_system_project_final_th_v2.tex"
AUX = SRC / "report_power_system_project_final_th_v2.aux"
FORM = ROOT / "docs" / "text" / "แบบฟอร์ม เอกสารสรุปวิชาโครงงาน.docx"
FIGDIR = SRC / "figures"
WORDFIG = FIGDIR / "final_report_th_v2" / "word"
OUT_DOCX = SRC / "report_power_system_project_final_th_v2.docx"
OUT_PDF = ROOT / "output" / "pdf" / "report_power_system_project_final_th_v2_word.pdf"

MML2OMML = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WNS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
THAI_FONT = "TH Sarabun New"
DOC_AUTHOR = "Phichitchai Jueajan"

# The form's own heading sizes: 18 pt for form-level headings, 16 pt for the
# rest.  Body size is resolved by the page-count search in main().
FORM_HEADING_PT = 18.0

# Cambria Math sits taller than TH Sarabun at the same nominal size, the same
# problem Computer Modern has against it in the PDF.  Measured digit heights:
# Sarabun 0.479 em, Cambria Math 0.674 em, Computer Modern 0.683 em.  The PDF
# solves it with \DeclareMathSizes at 0.84 of nominal; 0.86 of Cambria Math puts
# the digits at the same printed height, so the two documents look alike.
MATH_SCALE = 0.86

# Body type is the form's own size and is not a page-fitting lever.  The form's
# empty body paragraph marks carry sz=32, i.e. 16 pt.
BODY_PT = 16.0


@dataclass
class Layout:
    """Everything the page-count search is allowed to touch.

    Body type size, wording, headings, tables and figures all stay as they are;
    only the space around them and the raster scale move.
    """

    lead_pt: float          # exact line spacing for body text
    head_space: float       # multiplier on space above/below headings
    fig_scale: float        # multiplier on figure widths
    small_delta: float      # captions and table text, relative to body
    eq_space: float         # points above/below each equation line
    bib_columns: int        # bibliography column count


LAYOUT_LADDER = [
    Layout(17.6, 1.00, 1.00, -2.0, 3.0, 1),
    Layout(17.2, 0.85, 1.00, -2.0, 2.0, 1),
    Layout(16.8, 0.70, 1.00, -2.5, 2.0, 2),
    Layout(16.4, 0.60, 1.00, -3.0, 1.5, 2),
    Layout(16.0, 0.50, 1.00, -3.0, 1.0, 2),
    Layout(15.6, 0.42, 0.96, -3.0, 1.0, 2),
    Layout(15.2, 0.35, 0.92, -3.5, 0.5, 2),
    Layout(14.8, 0.30, 0.88, -3.5, 0.5, 2),
]

# Step 4 is the one that currently lands in range, so the search starts there
# and only moves if the measurement says otherwise.  Each trial costs a full
# Word round trip, so starting at the known answer usually means one.
LAYOUT_START = 4


# ---------------------------------------------------------------------------
# .aux reference-number map
# ---------------------------------------------------------------------------
def read_aux(path: Path) -> dict[str, str]:
    """Map each \\label to the number LaTeX printed for it.

    \\newlabel entries nest the printed form in braces, and the citation labels
    carry an extra pair because the bibliography prints [n] rather than n.  The
    innermost brace group is the printed number in both cases.
    """
    text = path.read_text(encoding="utf-8", errors="replace")
    out: dict[str, str] = {}
    for m in re.finditer(r"\\newlabel\{([^}]*)\}\{\{+([^{}]*)\}", text):
        out[m.group(1)] = m.group(2)
    if not out:
        raise SystemExit(f"no \\newlabel entries in {path}; run xelatex first")
    return out


# ---------------------------------------------------------------------------
# LaTeX maths -> OMML
# ---------------------------------------------------------------------------
class MathConverter:
    """LaTeX -> MathML (latex2mathml) -> OMML (Office's own MML2OMML.XSL).

    The XSL ships with Word and contains no msxsl:script block, so MSXML runs it
    without script permissions.  Word therefore receives native equations that
    stay editable rather than pictures of equations.

    Two MathML details need repair before the transform, both verified by
    inspecting the XSL's output directly:
      - latex2mathml emits \\dot and \\ddot as a plain <mover> without
        accent="true", which the XSL turns into a limit (the dot lands above and
        away from the letter) instead of an accent.  \\hat and \\overline already
        carry the attribute and come through correctly.
      - The XSL drops <mspace> entirely, so \\qquad and \\quad vanish and the two
        halves of an equation run together.  Substituting no-break spaces keeps
        the gap.
    Function names (\\max, \\sin, \\mathrm{sat}) arrive as <mo> and come out in
    italic; marking them as upright identifiers matches how LaTeX sets them.
    """

    # Accent characters latex2mathml places in a bare <mover>.
    _ACCENT_CHARS = {"˙", "¨", "^", "~", "´",
                     "`", "ˇ", "˘"}
    # Non-break spaces per mspace width, chosen so \qquad reads as a clear gap.
    _SPACE_WIDTH = {"2em": 6, "1em": 3, "0.5em": 2, "0.25em": 1,
                    "0.2em": 1, "0.167em": 1}
    # Operator names that must print upright, as they do in the PDF.
    _UPRIGHT = {"max", "min", "sin", "cos", "tan", "log", "ln", "exp",
                "lim", "sup", "inf", "det", "sat", "rcond", "Init", "diag"}
    # Bracket pairs that fence a matrix.
    _FENCES = {"[": "]", "(": ")", "{": "}", "|": "|",
               "‖": "‖", "⟨": "⟩"}
    _MML = "{http://www.w3.org/1998/Math/MathML}"

    def __init__(self) -> None:
        if not MML2OMML.exists():
            raise SystemExit(f"missing {MML2OMML}")
        self._xsl = win32com.client.Dispatch("Msxml2.DOMDocument.6.0")
        setattr(self._xsl, "async", False)
        if not self._xsl.load(str(MML2OMML)):
            raise SystemExit(f"cannot load {MML2OMML}: {self._xsl.parseError.reason}")
        self.count = 0

    @staticmethod
    def _upright_single(tex: str) -> str:
        r"""Rewrite one-character \mathrm{} as \text{}.

        latex2mathml only sets mathvariant="normal" when \mathrm wraps two or
        more characters, so \mathrm{s} in "s^{-1}" comes out italic while
        \mathrm{Hz} comes out upright.  \text{} takes the same path for both.
        """
        return re.sub(r"\\mathrm\{(.)\}", r"\\text{\1}", tex)

    @staticmethod
    def _brace_matrices(tex: str) -> str:
        """Brace a matrix that carries a trailing superscript.

        latex2mathml attaches the ^T to the opening bracket instead of the whole
        matrix, and the XSL then drops the closing bracket.  Wrapping the
        environment in a group makes both come out right.
        """
        for env in ("bmatrix", "pmatrix", "matrix", "vmatrix"):
            close = "\\end{" + env + "}"
            out, i = [], 0
            while True:
                j = tex.find(close, i)
                if j < 0:
                    out.append(tex[i:])
                    break
                k = j + len(close)
                if k < len(tex) and tex[k] in "^_":
                    start = tex.rfind("\\begin{" + env + "}", 0, j)
                    out.append(tex[i:start])
                    out.append("{" + tex[start:k] + "}")
                    i = k
                else:
                    out.append(tex[i:k])
                    i = k
            tex = "".join(out)
        return tex

    def _wrap_fences(self, root) -> None:
        """Group each bracket-matrix-bracket triple into its own <mrow>.

        latex2mathml leaves the fences as siblings of the table.  The XSL then
        pairs the first opening bracket with the *last* closing one in the whole
        expression, so a product of three matrices comes out as one giant
        bracket with stray brackets inside.  Wrapping each triple makes the XSL
        emit one m:d delimiter per matrix.
        """
        for parent in list(root.iter()):
            kids = list(parent)
            i = 0
            while i < len(kids) - 2:
                a, b, c = kids[i], kids[i + 1], kids[i + 2]
                if (a.tag == self._MML + "mo"
                        and (a.text or "") in self._FENCES
                        and b.tag == self._MML + "mtable"
                        and c.tag == self._MML + "mo"
                        and (c.text or "") == self._FENCES[a.text or ""]):
                    row = etree.Element(self._MML + "mrow")
                    parent.insert(list(parent).index(a), row)
                    for el in (a, b, c):
                        parent.remove(el)
                        row.append(el)
                    kids = list(parent)
                    i = 0
                    continue
                i += 1

    def _repair(self, mathml: str) -> str:
        root = etree.fromstring(mathml.encode("utf-8"))
        self._wrap_fences(root)
        for mover in root.iter(self._MML + "mover"):
            kids = list(mover)
            if (len(kids) == 2 and kids[1].tag == self._MML + "mo"
                    and (kids[1].text or "") in self._ACCENT_CHARS):
                mover.set("accent", "true")
        for op in root.iter(self._MML + "mo"):
            if (op.text or "").strip() in self._UPRIGHT:
                op.tag = self._MML + "mi"
                op.set("mathvariant", "normal")
        for space in list(root.iter(self._MML + "mspace")):
            n = self._SPACE_WIDTH.get(space.get("width", ""), 1)
            text = etree.Element(self._MML + "mtext")
            text.text = " " * n
            space.getparent().replace(space, text)
        return etree.tostring(root, encoding="unicode")

    def omml(self, tex: str, size_pt: float | None = None,
             display: bool = False) -> str:
        """Convert one maths fragment.  Raises on any failure: an equation that
        cannot be converted must stop the build, never vanish from the page.

        display=True asks for a display-mode zone, so fractions stack and sum
        limits sit above and below rather than beside the operator.
        """
        tex = tex.strip()
        mathml = self._repair(latex2mathml.converter.convert(
            self._upright_single(self._brace_matrices(tex)),
            display="block" if display else "inline"))
        dom = win32com.client.Dispatch("Msxml2.DOMDocument.6.0")
        setattr(dom, "async", False)
        dom.loadXML(mathml)
        if dom.parseError.errorCode != 0:
            raise ValueError(f"MathML rejected for {tex!r}: {dom.parseError.reason}")
        xml = dom.transformNode(self._xsl)
        if xml.startswith("<?xml"):
            xml = xml.split("?>", 1)[1]
        if "<m:oMath" not in xml:
            raise ValueError(f"no oMath produced for {tex!r}")
        if size_pt is not None:
            half = int(round(size_pt * 2))
            rpr = f'<w:rPr {WNS}><w:sz w:val="{half}"/><w:szCs w:val="{half}"/></w:rPr>'
            xml = xml.replace("<m:r>", "<m:r>" + rpr)
        self.count += 1
        return xml


# ---------------------------------------------------------------------------
# inline LaTeX -> a list of runs
# ---------------------------------------------------------------------------
@dataclass
class Frag:
    """One piece of a paragraph: literal text, or a maths fragment."""

    kind: str  # "text" | "math"
    text: str
    bold: bool = False
    italic: bool = False


# Greek and symbol names that appear in text mode via \emph{...} or headings.
TEXT_SUBST = {
    r"\GFL": "GFL",
    r"\leftrightarrow": "\u2194",
    r"\to": "\u2192",
    r"\second": " s",
    r"\pu": " pu",
    r"\hz": " Hz",
    r"\%": "%",
    r"\&": "&",
    r"\_": "_",
    r"\#": "#",
    r"\,": "\u2009",
    r"\ ": " ",
    r"\;": " ",
    r"\:": " ",
    r"~": "\u00a0",
    r"\@": "",
    r"\-": "",
    r"\par": "\n",
    "``": "\u201c",
    "''": "\u201d",
    "--": "\u2013",
}

# Accent/ligature forms used in the bibliography.
ACCENTS = {
    r"\~{n}": "\u00f1",
    r"\'{i}": "\u00ed",
    r"\'{c}": "\u0107",
    r"\'{\i}": "\u00ed",
}


def split_inline(src: str, macros: dict[str, str], refs: dict[str, str]) -> list[Frag]:
    """Split a LaTeX paragraph into text and maths fragments.

    Handles the subset the report actually uses (verified by enumerating every
    command in the body): \\emph, \\textbf, \\class, \\ref, \\eqref, the two
    \\input macro families, accents, and $...$ maths.
    """
    s = src
    for k, v in ACCENTS.items():
        s = s.replace(k, v)
    # A unit written as text plus a maths exponent (s$^{-1}$) leaves the letter
    # in Thai type and the exponent in Cambria, which reads as a stray digit.
    # Pull the letter into the same maths zone, upright as a unit should be.
    s = re.sub(r"(?<![A-Za-z$])([A-Za-z]{1,3})\$\^\{([^}]*)\}\$",
               lambda m: f"$\\mathrm{{{m.group(1)}}}^{{{m.group(2)}}}$", s)
    # \ref / \eqref resolve from the .aux
    def _ref(m: re.Match[str]) -> str:
        key = m.group(2)
        if key not in refs:
            raise ValueError(f"unresolved reference {key!r}; re-run xelatex")
        val = refs[key]
        return f"({val})" if m.group(1) == "eqref" else val

    s = re.sub(r"\\(eqref|ref)\{([^}]*)\}", _ref, s)
    # \input macros (\NewRun*, \Arm*, \Df*, \Set*)
    def _macro(m: re.Match[str]) -> str:
        name = m.group(1)
        if name in macros:
            return macros[name]
        return m.group(0)

    s = re.sub(r"\\([A-Za-z]+)(?![A-Za-z])", _macro, s)
    # drop label/index bookkeeping, then formatting-only commands
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    s = re.sub(r"\\(?:tablefont|tableheader|toprule|midrule|bottomrule|centering"
               r"|raggedright|arraybackslash|normalsize|small|scriptsize|bfseries"
               r"|noindent|newpage|hline)(?![A-Za-z])", "", s)
    s = re.sub(r"\\fontsize\{[^}]*\}\{[^}]*\}", "", s)
    s = re.sub(r"\\selectfont(?![A-Za-z])", "", s)
    s = re.sub(r"\\(?:v|h)space\*?\{[^}]*\}", "", s)

    frags: list[Frag] = []
    i, n = 0, len(s)
    buf: list[str] = []
    bold = italic = False

    def flush() -> None:
        if buf:
            frags.append(Frag("text", "".join(buf), bold, italic))
            buf.clear()

    while i < n:
        c = s[i]
        if c == "$":
            j = i + 1
            while j < n and not (s[j] == "$" and s[j - 1] != "\\"):
                j += 1
            flush()
            frags.append(Frag("math", s[i + 1 : j]))
            i = j + 1
            continue
        if c == "\\":
            m = re.match(r"\\(emph|textbf|textit|class)\{", s[i:])
            if m:
                depth, j = 1, i + m.end()
                while j < n and depth:
                    if s[j] == "{":
                        depth += 1
                    elif s[j] == "}":
                        depth -= 1
                    j += 1
                inner = s[i + m.end() : j - 1]
                flush()
                sub = split_inline(inner, macros, refs)
                for f in sub:
                    if f.kind == "text":
                        if m.group(1) in ("textbf", "class"):
                            f.bold = True
                        else:
                            f.italic = True
                    frags.append(f)
                i = j
                continue
            for k, v in TEXT_SUBST.items():
                if k.startswith("\\") and s.startswith(k, i):
                    nxt = i + len(k)
                    if k[-1].isalpha() and nxt < n and s[nxt].isalpha():
                        continue
                    buf.append(v)
                    i = nxt
                    break
            else:
                m2 = re.match(r"\\[A-Za-z]+\*?(\[[^\]]*\])?(\{[^{}]*\})?", s[i:])
                if m2:
                    raise ValueError(f"unhandled command {m2.group(0)!r} in {src[:70]!r}")
                buf.append(s[i + 1])
                i += 2
            continue
        for k, v in TEXT_SUBST.items():
            if not k.startswith("\\") and s.startswith(k, i):
                buf.append(v)
                i += len(k)
                break
        else:
            if c in "{}":
                i += 1
                continue
            buf.append(c)
            i += 1
    flush()
    out = [f for f in frags if f.kind == "math" or f.text]
    return out


# ---------------------------------------------------------------------------
# .tex -> ordered block list
# ---------------------------------------------------------------------------
@dataclass
class Block:
    kind: str  # formhead | formsub | section | subsection | para | result
               # | list | equation | table | figure
    text: str = ""
    items: list[str] = field(default_factory=list)
    lines: list[str] = field(default_factory=list)   # equation lines
    numbers: list[str] = field(default_factory=list) # equation numbers
    caption: str = ""
    label: str = ""
    rows: list[list[str]] = field(default_factory=list)
    aligns: list[str] = field(default_factory=list)
    widths: list[float] = field(default_factory=list)
    image: str = ""
    width_in: float = 0.0
    ordered: bool = True


def load_macros() -> dict[str, str]:
    """Read the two generated macro files.  Every printed scalar lives there."""
    out: dict[str, str] = {}
    for name in ("comparison_macros.tex", "run_summary_v2.tex"):
        p = FIGDIR / "switch_ieee14_decision" / name
        for m in re.finditer(r"\\newcommand\{\\([A-Za-z]+)\}\{(.*)\}\s*$",
                             p.read_text(encoding="utf-8"), re.M):
            out[m.group(1)] = m.group(2)
    return out


def strip_comments(text: str) -> str:
    return re.sub(r"(?<!\\)%.*", "", text)


def env_span(text: str, start: int, name: str) -> tuple[int, int]:
    """Return (inner_start, end_of_end_tag) for the environment opening at start."""
    open_tag = "\\begin{" + name + "}"
    close_tag = "\\end{" + name + "}"
    depth, i = 0, start
    while i < len(text):
        if text.startswith(open_tag, i):
            depth += 1
            i += len(open_tag)
            if depth == 1:
                inner = i
            continue
        if text.startswith(close_tag, i):
            depth -= 1
            i += len(close_tag)
            if depth == 0:
                return inner, i
            continue
        i += 1
    raise ValueError(f"unterminated {name}")


def parse_equation(body: str, env: str, star: bool, counter: list[int]) -> Block:
    """One display-maths block.

    Word has no amsmath, so every numbered line is emitted separately with its
    own number.  Row breaks and alignment markers are handled at environment
    depth zero only: a bmatrix carries its own \\\\ and & that must survive,
    while an align/aligned marker has no MathML equivalent (latex2mathml leaks a
    bare & that MSXML then rejects) and is dropped.
    """
    tag = re.search(r"\\tag\{([^}]*)\}", body)
    body = re.sub(r"\\tag\{[^}]*\}", "", body)
    body = re.sub(r"\\label\{[^}]*\}", "", body)
    inner = re.search(r"\\begin\{aligned\}(.*)\\end\{aligned\}", body, re.S)
    if inner:
        body = inner.group(1)
        one_number = True
    else:
        one_number = env == "equation"
    parts = [p.strip() for p in split_top(body, "\\\\") if p.strip()]
    lines = ["".join(split_top(p, "&")).strip() for p in parts]
    numbers: list[str] = []
    if star:
        numbers = [""] * len(lines)
    elif tag is not None:
        numbers = [tag.group(1)] + [""] * (len(lines) - 1)
    elif one_number:
        counter[0] += 1
        numbers = [str(counter[0])] + [""] * (len(lines) - 1)
    else:
        for _ in lines:
            counter[0] += 1
            numbers.append(str(counter[0]))
    return Block("equation", lines=lines, numbers=numbers)


def parse_table(body: str) -> Block:
    """One table.  Column alignment and relative widths come from the tabular
    preamble; l/c/r are sized from their content, L and p{..} take the rest."""
    cap = re.search(r"\\caption\{", body)
    caption = ""
    if cap:
        depth, i = 1, cap.end()
        while i < len(body) and depth:
            if body[i] == "{":
                depth += 1
            elif body[i] == "}":
                depth -= 1
            i += 1
        caption = body[cap.end() : i - 1]
    lab = re.search(r"\\label\{([^}]*)\}", body)
    label = lab.group(1) if lab else ""

    tm = re.search(r"\\begin\{tabularx?\}(?:\{[^{}]*(?:\{[^{}]*\})?[^{}]*\})?\{(.*?)\}\s*\n",
                   body, re.S)
    if not tm:
        raise ValueError("no tabular preamble found")
    spec = tm.group(1).replace("@{}", "")
    cols: list[tuple[str, float]] = []
    i = 0
    while i < len(spec):
        ch = spec[i]
        if ch in "lcr":
            cols.append(("lcr".index(ch) and ("center" if ch == "c" else "right")
                         or "left", 0.0))
            i += 1
        elif ch == "L" or ch == "X":
            cols.append(("left", -1.0))
            i += 1
        elif ch == "p" or ch == ">":
            m = re.match(r">\{[^{}]*(?:\{[^{}]*\})?[^{}]*\}", spec[i:])
            if m:
                i += m.end()
                continue
            m = re.match(r"p\{([0-9.]+)\\textwidth\}", spec[i:])
            if not m:
                raise ValueError(f"unparsed p column in {spec!r}")
            cols.append(("left", float(m.group(1))))
            i += m.end()
        elif ch in " \t\n":
            i += 1
        else:
            raise ValueError(f"unparsed column {ch!r} in {spec!r}")

    # Find the outer \end{tabular*} by depth, so a nested tabular inside a cell
    # (tables 5 and 9) does not terminate the scan early.
    inner_start = tm.end()
    depth, k, inner_end = 1, inner_start, -1
    while k < len(body):
        if body.startswith("\\begin{tabular", k):
            depth += 1
            k += 14
            continue
        if body.startswith("\\end{tabular", k):
            depth -= 1
            if depth == 0:
                inner_end = k
                break
            k += 12
            continue
        k += 1
    if inner_end < 0:
        raise ValueError("unterminated tabular")
    inner = body[inner_start:inner_end]
    inner = re.sub(r"\\(?:toprule|midrule|bottomrule)", "", inner)
    rows: list[list[str]] = []
    for raw in split_top(inner, "\\\\"):
        raw = raw.strip()
        if not raw:
            continue
        cells = split_top(raw, "&")
        if len(cells) != len(cols):
            raise ValueError(f"row has {len(cells)} cells, spec has {len(cols)}: {raw[:80]!r}")
        rows.append([c.strip() for c in cells])
    return Block("table", caption=caption, label=label, rows=rows,
                 aligns=[a for a, _ in cols], widths=[w for _, w in cols])


def split_top(src: str, sep: str) -> list[str]:
    """Split on sep at brace depth zero and outside any nested environment.

    Both guards are needed: cells hold braced groups, and tables 5 and 9 nest a
    tabular whose own row breaks must not split the outer row.
    """
    out: list[str] = []
    buf: list[str] = []
    depth = env = 0
    i, n = 0, len(src)
    while i < n:
        m = re.match(r"\\(begin|end)\{[^}]*\}", src[i:])
        if m:
            env += 1 if m.group(1) == "begin" else -1
            buf.append(m.group(0))
            i += m.end()
            continue
        if depth == 0 and env == 0 and src.startswith(sep, i):
            out.append("".join(buf))
            buf = []
            i += len(sep)
            continue
        c = src[i]
        if c == "\\" and i + 1 < n and not src[i + 1].isalpha():
            buf.append(src[i : i + 2])
            i += 2
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
        buf.append(c)
        i += 1
    out.append("".join(buf))
    return out


# Headings: the form's exact strings, in the form's order.  The .tex uses a
# different wording for the second one; the form wins, as the owner directed.
FORM_HEADINGS = [
    "วัตถุประสงค์ของโครงงาน",
    "ความสำคัญ และที่มาของปัญหาที่ทำโครงงาน",
    "การดำเนินงานและผลงานที่ได้รับจากโครงงาน (โดยสังเขป) พร้อมภาพประกอบ",
    "ประโยชน์ที่ได้รับจากโครงงาน",
    "ปัญหาและอุปสรรคที่เกิดขึ้น",
    "แนวทางการแก้ไขปัญหาและอุปสรรคที่เกิดขึ้น",
    "โครงงานที่คาดว่าจะดำเนินการต่อไป",
    "เอกสารอ้างอิง",
]

# Figure widths in inches inside the form's 5.7681 in column.
FIG_WIDTH = {
    "fig:network": 3.70,
    "fig:pipeline": 5.10,
    "fig:dualibr": 5.7681,
    "fig:supervisor": 5.7681,
    "fig:periods": 5.00,
    "fig:electrical": 5.7681,
    "fig:bus9": 5.7681,
}


def parse_tex(text: str, macros: dict[str, str]) -> list[Block]:
    """Walk the document body in order and emit a flat block list."""
    body = strip_comments(text.split("\\begin{document}", 1)[1])
    body = body.split("\\end{titlepage}", 1)[1]
    body = body.split("\\end{document}", 1)[0]

    blocks: list[Block] = []
    eq_counter = [0]
    sec_no = [0, 0]
    form_no = [0]
    pending_runhead: list[str] = []
    i, n = 0, len(body)
    buf: list[str] = []

    def flush_para() -> None:
        txt = "".join(buf).strip()
        buf.clear()
        chunks = [" ".join(c.split()) for c in re.split(r"\n\s*\n", txt)]
        chunks = [c for c in chunks if c]
        if pending_runhead and chunks:
            blocks.append(Block("runhead", text=pending_runhead[0],
                                items=[chunks[0]]))
            chunks = chunks[1:]
            pending_runhead.clear()
        for chunk in chunks:
            blocks.append(Block("para", text=chunk))

    while i < n:
        if body.startswith("\\templateheading{", i):
            flush_para()
            j = body.index("}", i)
            form_no[0] += 1
            blocks.append(Block("formhead", text=FORM_HEADINGS[form_no[0] - 1]))
            i = j + 1
            continue
        if body.startswith("\\templatesubheading{", i):
            flush_para()
            depth, j = 1, i + len("\\templatesubheading{")
            while depth:
                if body[j] == "{":
                    depth += 1
                elif body[j] == "}":
                    depth -= 1
                j += 1
            blocks.append(Block("formsub", text=body[i + len("\\templatesubheading{") : j - 1]))
            i = j
            continue
        m = re.match(r"\\(sub)?section\{", body[i:])
        if m:
            flush_para()
            depth, j = 1, i + m.end()
            while depth:
                if body[j] == "{":
                    depth += 1
                elif body[j] == "}":
                    depth -= 1
                j += 1
            title = body[i + m.end() : j - 1]
            if m.group(1):
                sec_no[1] += 1
                blocks.append(Block("subsection",
                                    text=f"{sec_no[0]}.{sec_no[1]} {title}"))
            else:
                sec_no[0] += 1
                sec_no[1] = 0
                blocks.append(Block("section", text=f"{sec_no[0]}. {title}"))
            i = j
            continue
        if body.startswith("\\paragraph{", i):
            flush_para()
            depth, j = 1, i + len("\\paragraph{")
            while depth:
                if body[j] == "{":
                    depth += 1
                elif body[j] == "}":
                    depth -= 1
                j += 1
            pending_runhead.append(body[i + len("\\paragraph{") : j - 1])
            i = j
            continue
        if body.startswith("\\resultbox{", i):
            flush_para()
            depth, j = 1, i + len("\\resultbox{")
            while depth:
                if body[j] == "{":
                    depth += 1
                elif body[j] == "}":
                    depth -= 1
                j += 1
            blocks.append(Block("result", text=" ".join(
                body[i + len("\\resultbox{") : j - 1].split())))
            i = j
            continue
        m = re.match(r"\\begin\{(equation|align|gather|multline)(\*?)\}", body[i:])
        if m:
            flush_para()
            env, star = m.group(1), m.group(2) == "*"
            name = env + m.group(2)
            inner, end = env_span(body, i, name)
            blocks.append(parse_equation(body[inner : end - len("\\end{" + name + "}")],
                                         env, star, eq_counter))
            i = end
            continue
        if body.startswith("\\begin{table}", i):
            flush_para()
            inner, end = env_span(body, i, "table")
            blocks.append(parse_table(body[inner:end]))
            i = end
            continue
        if body.startswith("\\begin{figure}", i):
            flush_para()
            inner, end = env_span(body, i, "figure")
            blocks.append(parse_figure(body[inner:end], len(blocks)))
            i = end
            continue
        if body.startswith("\\begin{enumerate}", i) or body.startswith("\\begin{itemize}", i):
            flush_para()
            name = "enumerate" if body.startswith("\\begin{enumerate}", i) else "itemize"
            inner, end = env_span(body, i, name)
            raw = body[inner : end - len("\\end{" + name + "}")]
            raw = re.sub(r"^\s*\[[^\]]*\]", "", raw)
            items = [" ".join(x.split()) for x in re.split(r"\\item", raw)[1:]]
            blocks.append(Block("list", items=[x for x in items if x],
                                ordered=(name == "enumerate")))
            i = end
            continue
        if body.startswith("\\begin{multicols}", i):
            flush_para()
            inner, end = env_span(body, i, "multicols")
            raw = body[inner : end - len("\\end{multicols}")]
            inner2, end2 = env_span(raw, raw.index("\\begin{enumerate}"), "enumerate")
            items_raw = raw[inner2 : end2 - len("\\end{enumerate}")]
            items_raw = re.sub(r"^\s*\[[^\]]*\]", "", items_raw)
            items = [" ".join(x.split()) for x in re.split(r"\\item", items_raw)[1:]]
            blocks.append(Block("biblist", items=[x for x in items if x]))
            i = end
            continue
        if body.startswith("\\newpage", i):
            i += len("\\newpage")
            continue
        buf.append(body[i])
        i += 1
    flush_para()
    return blocks


def parse_figure(body: str, order: int) -> Block:
    """One figure.  A TikZ body is rendered to PNG; \\includegraphics is used
    as-is at the width the form's column allows."""
    cap = re.search(r"\\caption\{", body)
    caption = ""
    if cap:
        depth, i = 1, cap.end()
        while i < len(body) and depth:
            if body[i] == "{":
                depth += 1
            elif body[i] == "}":
                depth -= 1
            i += 1
        caption = " ".join(body[cap.end() : i - 1].split())
    lab = re.search(r"\\label\{([^}]*)\}", body)
    label = lab.group(1) if lab else ""
    inc = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}", body)
    if inc:
        name = inc.group(1)
        for cand in (FIGDIR / name,
                     FIGDIR / "final_report_th_v2" / name,
                     FIGDIR / "presentation_pf_sssa_ts_gfm" / name):
            if cand.exists():
                image = str(cand)
                break
        else:
            raise SystemExit(f"figure not found: {name}")
    else:
        image = str(WORDFIG / f"{label.replace(':', '_')}.png")
    return Block("figure", caption=caption, label=label, image=image,
                 width_in=FIG_WIDTH[label])


# ---------------------------------------------------------------------------
# TikZ -> PNG, rendered from the report's own source
# ---------------------------------------------------------------------------
TIKZ_PREAMBLE = r"""\documentclass[border=3pt]{standalone}
\usepackage{fontspec}
\XeTeXlinebreaklocale "th"
\XeTeXlinebreakskip=0pt plus 1pt
\defaultfontfeatures{Ligatures=TeX}
\setmainfont[
  Path={%(fontpath)s},
  Extension=.ttf,
  UprightFont={*}, BoldFont={* Bold},
  ItalicFont={* Italic}, BoldItalicFont={* BoldItalic}
]{THSarabun}
\usepackage{type1cm}
\usepackage{amsmath,amssymb,mathtools}
\usepackage[table]{xcolor}
\usepackage{tikz}
\usetikzlibrary{arrows.meta,positioning,calc,shapes.geometric}
\definecolor{academicblue}{RGB}{35,64,96}
\definecolor{blockfill}{RGB}{242,244,247}
\newcommand{\second}{\ensuremath{\,\mathrm{s}}}
\input{%(macros)s}
\begin{document}\fontsize{15}{16.5}\selectfont
"""


def render_tikz(text: str, force: bool = False) -> None:
    """Render the pipeline and chronology diagrams from the .tex to PNG.

    The diagrams stay in LaTeX source, so the Thai lettering and the
    \\NewRun* values in them are the same ones the PDF prints.
    """
    body = strip_comments(text)
    spans = [m.start() for m in re.finditer(r"\\begin\{tikzpicture\}", body)]
    # spans[0] is the title-page rule; the two diagrams follow it.
    targets = {"fig_pipeline": spans[1], "fig_periods": spans[2]}
    WORDFIG.mkdir(parents=True, exist_ok=True)
    fontpath = (SRC / "fonts" / "TH Sarabun PSK V-1" / "TH Sarabun PSK V-1")
    head = TIKZ_PREAMBLE % {
        "fontpath": str(fontpath).replace("\\", "/") + "/",
        "macros": str(FIGDIR / "switch_ieee14_decision" / "run_summary_v2.tex").replace("\\", "/"),
    }
    for name, start in targets.items():
        png = WORDFIG / f"{name}.png"
        if png.exists() and not force:
            continue
        inner, end = env_span(body, start, "tikzpicture")
        snippet = body[start:end]
        work = WORDFIG / "_tikz"
        work.mkdir(exist_ok=True)
        tex = work / f"{name}.tex"
        tex.write_text(head + snippet + "\n\\end{document}\n", encoding="utf-8")
        r = subprocess.run(["xelatex", "-interaction=nonstopmode", tex.name],
                           cwd=work, capture_output=True)
        pdf = work / f"{name}.pdf"
        if r.returncode != 0 or not pdf.exists():
            log = (work / f"{name}.log").read_text(encoding="utf-8", errors="replace")
            errs = [l for l in log.splitlines() if l.startswith("! ")]
            raise SystemExit(f"TikZ render failed for {name}: {errs[:3]}")
        subprocess.run(["pdftoppm", "-r", "400", "-png", "-singlefile",
                        str(pdf), str(png.with_suffix(""))], check=True)
        print(f"  rendered {png.name}")


# ---------------------------------------------------------------------------
# DOCX emission
# ---------------------------------------------------------------------------
class Builder:
    def __init__(self, layout: Layout, math: MathConverter,
                 macros: dict[str, str], refs: dict[str, str]) -> None:
        self.L = layout
        self.body_pt = BODY_PT
        self.small_pt = BODY_PT + layout.small_delta
        self.math = math
        self.macros = macros
        self.refs = refs
        shutil.copy2(FORM, OUT_DOCX)
        self.doc = docx.Document(str(OUT_DOCX))
        el = self.doc.element.body
        for child in list(el):
            if not child.tag.endswith("sectPr"):
                el.remove(child)
        sec = self.doc.sections[0]
        self.text_in = (int(sec.page_width) - int(sec.left_margin)
                        - int(sec.right_margin)) / 914400.0
        self._set_default_font()
        self.fig_no = 0
        self.tab_no = 0

    def _set_default_font(self) -> None:
        """Point every style in the form at TH Sarabun New.

        The form's docDefaults name Times New Roman with Angsana New for complex
        script, Normal names Cordia New, and the heading styles name Arial and
        Cambria.  Any run this script does not style explicitly - a paragraph
        mark, an empty table cell, a footer field - would otherwise embed those
        faces in the file and in the PDF.  Only Cambria Math is left alone: it
        is the maths font and Word requires it there.
        """
        rpr = (f'<w:rPr {WNS}><w:rFonts w:ascii="{THAI_FONT}" '
               f'w:hAnsi="{THAI_FONT}" w:cs="{THAI_FONT}" '
               f'w:eastAsia="{THAI_FONT}"/>'
               f'<w:sz w:val="{int(BODY_PT * 2)}"/>'
               f'<w:szCs w:val="{int(BODY_PT * 2)}"/></w:rPr>')
        styles = self.doc.styles.element
        defaults = styles.find(qn("w:docDefaults"))
        rpr_default = defaults.find(qn("w:rPrDefault"))
        for old in rpr_default.findall(qn("w:rPr")):
            rpr_default.remove(old)
        rpr_default.append(parse_xml(rpr))
        # Rewrite the font name on every w:rFonts in every style definition.
        for fonts in styles.iter(qn("w:rFonts")):
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                if fonts.get(qn("w:" + attr)) not in (None, "Cambria Math"):
                    fonts.set(qn("w:" + attr), THAI_FONT)
            for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
                if fonts.get(qn("w:" + attr)) is not None:
                    del fonts.attrib[qn("w:" + attr)]

    # -- low-level -----------------------------------------------------------
    def _style_run(self, run, size: float, bold: bool) -> None:
        run.font.size = Pt(size)
        run.font.name = THAI_FONT
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = parse_xml(f"<w:rFonts {WNS}/>")
            rpr.insert(0, rf)
        for attr in ("ascii", "hAnsi", "cs"):
            rf.set(qn("w:" + attr), THAI_FONT)
        half = int(round(size * 2))
        rpr.append(parse_xml(f'<w:szCs {WNS} w:val="{half}"/>'))
        # Mark the run as complex script.  Without this Word justifies Thai as
        # if it were Latin and spreads the letters inside words; the form's own
        # paragraphs carry the same flag.
        rpr.append(parse_xml(f"<w:cs {WNS}/>"))
        rpr.append(parse_xml(f'<w:lang {WNS} w:val="en-US" w:bidi="th-TH"/>'))
        if bold:
            run.bold = True
            rpr.append(parse_xml(f"<w:bCs {WNS}/>"))

    def _para(self, size: float | None = None, align=None, indent: float = 0.0,
              before: float = 0.0, after: float = 0.0, lead: float | None = None,
              keep: bool = False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = Pt(self.L.lead_pt) if lead is None else lead
        if indent:
            pf.first_line_indent = Inches(indent)
        if align is not None:
            p.alignment = align
        if keep:
            pf.keep_with_next = True
        return p

    def _emit_frags(self, p, frags: list[Frag], size: float,
                    bold: bool = False) -> None:
        for f in frags:
            if f.kind == "math":
                p._p.append(parse_xml(self.math.omml(f.text, size * MATH_SCALE)))
            else:
                run = p.add_run(f.text)
                self._style_run(run, size, f.bold or bold)
                if f.italic:
                    run.italic = True

    def _text(self, p, src: str, size: float, bold: bool = False) -> None:
        self._emit_frags(p, split_inline(src, self.macros, self.refs), size, bold)

    def _thai_justify(self, p) -> None:
        """Start each line flush left and let Word distribute the slack the Thai
        way, as the owner asked and as the form's own body paragraphs are set.

        Thai has no hyphenation, so `both` has to stretch something to reach the
        right margin; thaiDistribute spreads the line by its own rules instead
        of pulling Latin words apart.
        """
        ppr = p._p.get_or_add_pPr()
        for old in ppr.findall(qn("w:jc")):
            ppr.remove(old)
        ppr.append(parse_xml(f'<w:jc {WNS} w:val="thaiDistribute"/>'))

    # -- cover ---------------------------------------------------------------
    def cover(self) -> None:
        """The form's cover fields, filled with this project's values."""
        C = WD_ALIGN_PARAGRAPH.CENTER
        p = self._para(align=C, before=6, after=18, lead=1.0)
        self._style_run(p.add_run("เอกสารสรุปวิชาโครงงาน"), 20, True)
        rows = [
            ("ชื่องานโครงงาน", 18, True, 4, 2),
            ("(ภาษาไทย) กรอบการทำงานสำหรับการสลับโหมด Grid-Following/Grid-Forming"
             " แบบอัตโนมัติของ Inverter-Based Resources", 16, False, 0, 2),
            ("(ภาษาอังกฤษ) An Automatic Following/Forming Framework"
             " for Inverter-Based Resources", 16, False, 0, 14),
            ("สมาชิกผู้จัดทำ", 18, True, 0, 2),
            ("1. พิชิตชัยเจือจันทร์\t66010569", 16, False, 0, 0),
            ("2. ภัคพล ผดุงแดน\t66010612", 16, False, 0, 14),
            ("อาจารย์ที่ปรึกษา", 18, True, 0, 2),
            ("1. ดร. ทศพร สุรินทร์แก้ว", 16, False, 0, 0),
            ("2. ศ. ดร. อิสระชัย งามหรู", 16, False, 0, 24),
            ("ภาควิชาวิศวกรรมไฟฟ้า คณะวิศวกรรมศาสตร์", 16, False, 0, 0),
            ("สถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง", 16, False, 0, 12),
            ("ปีการศึกษา 2569", 18, True, 0, 0),
        ]
        for text, size, bold, before, after in rows:
            p = self._para(align=C, before=before, after=after, lead=1.0)
            if "\t" in text:
                p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(self.text_in * 0.62), WD_TAB_ALIGNMENT.LEFT)
            self._style_run(p.add_run(text), size, bold)
        self.doc.add_page_break()

    # -- blocks --------------------------------------------------------------
    def formhead(self, text: str) -> None:
        h = self.L.head_space
        p = self._para(before=10 * h, after=4 * h, keep=True, lead=1.0)
        self._style_run(p.add_run(text), FORM_HEADING_PT, True)

    def formsub(self, text: str) -> None:
        h = self.L.head_space
        p = self._para(before=7 * h, after=3 * h, keep=True, lead=1.0)
        self._style_run(p.add_run(text), self.body_pt + 1.0, True)

    def section(self, text: str) -> None:
        h = self.L.head_space
        p = self._para(before=8 * h, after=3 * h, keep=True, lead=1.0)
        self._text(p, text, self.body_pt + 1.0, bold=True)

    def subsection(self, text: str) -> None:
        h = self.L.head_space
        p = self._para(before=6 * h, after=2 * h, keep=True, lead=1.0)
        self._text(p, text, self.body_pt + 0.5, bold=True)

    def para(self, text: str) -> None:
        p = self._para(indent=0.25)
        self._thai_justify(p)
        self._text(p, text, self.body_pt)

    def runhead(self, blk: Block) -> None:
        """\\paragraph: a bold run-in lead-in followed by its own text, on one
        line, the way LaTeX sets it."""
        p = self._para(before=3, indent=0.25)
        self._thai_justify(p)
        self._style_run(p.add_run(blk.text + "  "), self.body_pt, True)
        if blk.items:
            self._text(p, blk.items[0], self.body_pt)

    def result(self, text: str) -> None:
        p = self._para(before=3, after=3, indent=0.25)
        self._thai_justify(p)
        self._style_run(p.add_run("ผลสำคัญ. "), self.body_pt, True)
        self._text(p, text, self.body_pt)

    def lst(self, items: list[str], ordered: bool) -> None:
        for k, item in enumerate(items, 1):
            p = self._para()
            self._thai_justify(p)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.36)
            pf.first_line_indent = Inches(-0.36)
            pf.tab_stops.add_tab_stop(Inches(0.36), WD_TAB_ALIGNMENT.LEFT)
            self._style_run(p.add_run(f"{k}.\t" if ordered else "•\t"),
                            self.body_pt, False)
            self._text(p, item, self.body_pt)

    def biblist(self, items: list[str]) -> None:
        size = self.small_pt + 1.0
        if self.L.bib_columns > 1:
            self._set_columns(self.L.bib_columns)
        for k, item in enumerate(items, 1):
            p = self._para(lead=1.0)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.36)
            pf.first_line_indent = Inches(-0.36)
            pf.tab_stops.add_tab_stop(Inches(0.36), WD_TAB_ALIGNMENT.LEFT)
            self._style_run(p.add_run(f"[{k}]\t"), size, False)
            self._text(p, item, size)

    def _set_columns(self, n: int) -> None:
        """Put the bibliography in n columns, as the PDF does with multicols.

        A continuous section break carries the column count; the surrounding
        single-column layout is unaffected because this is the last block.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        ppr = p._p.get_or_add_pPr()
        ppr.append(parse_xml(
            f'<w:sectPr {WNS}><w:type w:val="continuous"/>'
            f'<w:pgSz w:w="11906" w:h="16838"/>'
            f'<w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800"'
            f' w:header="720" w:footer="720" w:gutter="0"/>'
            f'<w:cols w:num="1" w:space="720"/></w:sectPr>'))
        sec = self.doc.sections[-1]
        sec_pr = sec._sectPr
        cols = sec_pr.find(qn("w:cols"))
        cols.set(qn("w:num"), str(n))
        cols.set(qn("w:space"), "360")

    def equation(self, lines: list[str], numbers: list[str]) -> None:
        """A numbered display equation.

        The maths and its number go in a borderless two-cell table.  A tab-and-
        run layout was tried first and Word demoted the equation to inline size
        (fractions squeezed onto one line, sum limits pushed to the side); a
        table cell keeps the equation a display-mode zone at full size while the
        number stays flush right, which is how the PDF sets it.
        """
        for line, num in zip(lines, numbers):
            t = self.doc.add_table(rows=1, cols=2)
            t.autofit = False
            t._tbl.tblPr.append(parse_xml(
                f'<w:tblBorders {WNS}><w:top w:val="none"/><w:left w:val="none"/>'
                f'<w:bottom w:val="none"/><w:right w:val="none"/>'
                f'<w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>'))
            num_w = 0.62
            cell_eq, cell_no = t.rows[0].cells
            cell_eq.width = Inches(self.text_in - num_w)
            cell_no.width = Inches(num_w)
            p = cell_eq.paragraphs[0]
            p.paragraph_format.space_before = Pt(self.L.eq_space)
            p.paragraph_format.space_after = Pt(self.L.eq_space)
            p.paragraph_format.line_spacing = 1.0
            p._p.append(parse_xml(
                self.math.omml(line, self.body_pt * MATH_SCALE, display=True)))
            q = cell_no.paragraphs[0]
            q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            q.paragraph_format.space_before = Pt(self.L.eq_space)
            q.paragraph_format.space_after = Pt(self.L.eq_space)
            q.paragraph_format.line_spacing = 1.0
            cell_no.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if num:
                self._style_run(q.add_run(f"({num})"), self.small_pt + 1.0, False)

    def table(self, blk: Block) -> None:
        self.tab_no += 1
        size = self.small_pt
        cap = self._para(align=WD_ALIGN_PARAGRAPH.CENTER,
                         before=5 * self.L.head_space, after=2, lead=1.0,
                         keep=True)
        self._style_run(cap.add_run(f"ตารางที่ {self.tab_no} "), size, True)
        self._text(cap, blk.caption, size)

        ncol = len(blk.aligns)
        widths = self._col_widths(blk, size)
        t = self.doc.add_table(rows=len(blk.rows), cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        amap = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT}
        for ri, row in enumerate(blk.rows):
            tr = t.rows[ri]
            if ri == 0:
                tr._tr.get_or_add_trPr().append(parse_xml(f"<w:tblHeader {WNS}/>"))
            for ci, cell in enumerate(row):
                tc = tr.cells[ci]
                tc.width = Inches(widths[ci])
                header = ri == 0
                self._fill_cell(tc, cell, size, amap[blk.aligns[ci]], header)
        self._para(before=0, after=3, lead=1.0)

    def _fill_cell(self, tc, src: str, size: float, align, header: bool) -> None:
        """Fill one cell.  A nested tabular becomes one line per row so the
        stacked state lists in tables 5 and 9 keep their layout."""
        nested = re.search(
            r"\\begin\{tabular\}(?:\[[^\]]*\])?"
            r"\{(?:[^{}]|\{[^{}]*\})*\}(.*?)\\end\{tabular\}", src, re.S)
        parts = [src]
        if nested:
            parts = [x.strip() for x in split_top(nested.group(1), "\\\\") if x.strip()]
        first = True
        for part in parts:
            p = tc.paragraphs[0] if first else tc.add_paragraph()
            first = False
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            p.alignment = align
            frags = split_inline(part, self.macros, self.refs)
            for f in frags:
                if f.kind == "math":
                    p._p.append(parse_xml(
                        self.math.omml(f.text, size * MATH_SCALE)))
                else:
                    run = p.add_run(f.text)
                    self._style_run(run, size, f.bold or header)
                    if f.italic:
                        run.italic = True
            if all(f.kind == "math" for f in frags):
                # Word centres a paragraph that holds nothing but an equation,
                # ignoring w:jc.  One zero-width space restores the cell's own
                # alignment without printing anything.
                self._style_run(p.add_run("​"), size, False)

    def _col_widths(self, blk: Block, size: float) -> list[float]:
        """Size l/c/r columns from their widest cell, give p{} its declared
        fraction, and split what remains between the L columns."""
        from fontTools.ttLib import TTFont

        if not hasattr(Builder, "_metrics"):
            font = TTFont(os.path.join(os.environ["LOCALAPPDATA"],
                                       r"Microsoft\Windows\Fonts\THSarabunNew.ttf"))
            math_font = TTFont(r"C:\Windows\Fonts\cambria.ttc", fontNumber=1)
            Builder._metrics = (font.getBestCmap(), font["hmtx"],
                                font["head"].unitsPerEm,
                                math_font.getBestCmap(), math_font["hmtx"],
                                math_font["head"].unitsPerEm)
        cmap, hmtx, upem, mcmap, mhmtx, mupem = Builder._metrics

        def math_width(tex: str) -> float:
            """Printed width of a maths fragment in em, from Cambria Math's own
            advance widths.  Commands stand in for one glyph, braces for none,
            and super/subscripts are counted at 0.65 of full size."""
            def adv(s: str, scale: float = 1.0) -> float:
                t = 0.0
                for ch in s:
                    g = mcmap.get(ord(ch))
                    if g is not None:
                        t += mhmtx[g][0]
                return t / mupem * scale

            total = 0.0
            i = 0
            while i < len(tex):
                m = re.match(r"[\^_]\{([^{}]*)\}|[\^_](.)", tex[i:])
                if m:
                    total += adv(m.group(1) or m.group(2) or "", 0.65)
                    i += m.end()
                    continue
                m = re.match(r"\\times", tex[i:])
                if m:
                    total += adv("×")
                    i += m.end()
                    continue
                m = re.match(r"\\[A-Za-z]+", tex[i:])
                if m:
                    total += adv("m")  # a Greek letter or operator: one glyph
                    i += m.end()
                    continue
                ch = tex[i]
                i += 1
                if ch in "{}\\ \t":
                    continue
                total += adv(ch)
            return total

        def measure(text: str) -> float:
            frags = split_inline(text, self.macros, self.refs)
            em = 0.0
            for f in frags:
                if f.kind == "math":
                    em += math_width(f.text) * MATH_SCALE
                else:
                    for ch in f.text:
                        g = cmap.get(ord(ch))
                        if g is not None:
                            em += hmtx[g][0] / upem
            return em * size / 72.0

        pad = 2 * 0.055  # Word's default cell margins, both sides
        avail = self.text_in
        widths = [0.0] * len(blk.aligns)
        flex: list[int] = []
        for ci, w in enumerate(blk.widths):
            if w < 0:
                flex.append(ci)
            elif w > 0:
                widths[ci] = w * self.text_in
            else:
                widths[ci] = min(max(measure(r[ci]) for r in blk.rows) + pad,
                                 self.text_in * 0.32)
        used = sum(widths)
        if flex:
            # Every flex column needs room for a few Thai words per line.
            share = (avail - used) / len(flex)
            if share < 1.10:
                shrink = min(1.0, max(0.0, (avail - 1.10 * len(flex)) / used)) if used else 1.0
                widths = [w * shrink for w in widths]
                used = sum(widths)
                share = (avail - used) / len(flex)
            for ci in flex:
                widths[ci] = share
        total = sum(widths)
        if total > avail:
            widths = [w * avail / total for w in widths]
        elif total < avail * 0.995:
            # Spread the slack over the text columns so the table fills the
            # column the way the PDF's tabularx does, instead of sitting narrow.
            text_cols = [ci for ci, a in enumerate(blk.aligns) if a == "left"]
            targets = text_cols or list(range(len(widths)))
            extra = (avail - total) / len(targets)
            for ci in targets:
                widths[ci] += extra
        return widths

    def figure(self, blk: Block) -> None:
        self.fig_no += 1
        width = min(blk.width_in * self.L.fig_scale, self.text_in)
        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER,
                       before=5 * self.L.head_space, after=2, lead=1.0,
                       keep=True)
        p.add_run().add_picture(blk.image, width=Inches(width))
        size = self.small_pt
        cap = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4,
                         lead=1.0)
        self._thai_justify(cap)
        cap.paragraph_format.left_indent = Inches(0.25)
        cap.paragraph_format.right_indent = Inches(0.25)
        self._style_run(cap.add_run(f"รูปที่ {self.fig_no} "), size, True)
        self._text(cap, blk.caption, size)

    # -- page numbers --------------------------------------------------------
    def fix_footer(self) -> None:
        """Give the form's page number a footer that actually renders.

        As distributed, sectPr references a footer of type "even" while
        settings.xml omits evenAndOddHeaders, so Word shows no page number on
        any page.  A default footer carrying the same PAGE field, in the form's
        own Footer and PageNumber styles, makes it appear on every page.
        """
        sec = self.doc.sections[0]
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.style = self.doc.styles["Footer"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        self._style_run(run, self.small_pt, False)
        rpr = run._element.get_or_add_rPr()
        for xml in (f'<w:fldChar {WNS} w:fldCharType="begin"/>',
                    f'<w:instrText {WNS} xml:space="preserve">PAGE</w:instrText>',
                    f'<w:fldChar {WNS} w:fldCharType="end"/>'):
            run._element.append(parse_xml(xml))
        sec.footer_distance = Inches(0.40)

    def properties(self) -> None:
        """Set the file's own properties.

        The form arrives with dc:title "1", dc:creator "Power" and a created
        date of 2020-10-02 from whoever last saved it; all three would otherwise
        ship inside the deliverable.
        """
        props = self.doc.core_properties
        props.title = ("กรอบการทำงานสำหรับการสลับโหมด Grid-Following/Grid-Forming"
                       " แบบอัตโนมัติของ Inverter-Based Resources")
        props.subject = "เอกสารสรุปวิชาโครงงาน"
        props.author = DOC_AUTHOR
        props.last_modified_by = DOC_AUTHOR
        props.category = ""
        props.comments = ""
        props.keywords = ""
        props.revision = 1
        now = datetime.now().replace(microsecond=0)
        props.created = now
        props.modified = now

    # -- driver --------------------------------------------------------------
    def build(self, blocks: list[Block]) -> None:
        self.cover()
        dispatch = {
            "formhead": lambda b: self.formhead(b.text),
            "formsub": lambda b: self.formsub(b.text),
            "section": lambda b: self.section(b.text),
            "subsection": lambda b: self.subsection(b.text),
            "runhead": self.runhead,
            "para": lambda b: self.para(b.text),
            "result": lambda b: self.result(b.text),
            "list": lambda b: self.lst(b.items, b.ordered),
            "biblist": lambda b: self.biblist(b.items),
            "equation": lambda b: self.equation(b.lines, b.numbers),
            "table": self.table,
            "figure": self.figure,
        }
        for blk in blocks:
            dispatch[blk.kind](blk)
        self.fix_footer()
        self.properties()
        self.doc.save(str(OUT_DOCX))


# ---------------------------------------------------------------------------
# Word round-trip: page count, PDF export, structural gates
# ---------------------------------------------------------------------------
def word_measure(path: Path, pdf: Path | None = None) -> tuple[int, int]:
    """Open in Word, return (pages, equation count), optionally export PDF."""
    app = win32com.client.DispatchEx("Word.Application")
    app.Visible = False
    app.DisplayAlerts = 0
    try:
        doc = app.Documents.Open(str(path), ReadOnly=False, AddToRecentFiles=False)
        try:
            doc.Fields.Update()
            pages = doc.ComputeStatistics(2)  # wdStatisticPages
            omaths = doc.OMaths.Count
            # Word stamps whoever is signed in as the last author on save; put
            # the project's own name back before the PDF is produced from it.
            doc.BuiltInDocumentProperties("Last author").Value = DOC_AUTHOR
            doc.BuiltInDocumentProperties("Author").Value = DOC_AUTHOR
            if pdf is not None:
                pdf.parent.mkdir(parents=True, exist_ok=True)
                doc.ExportAsFixedFormat(str(pdf), 17)  # wdExportFormatPDF
            doc.Save()
        finally:
            doc.Close(False)
    finally:
        app.Quit()
    return pages, omaths


def check_package(path: Path) -> None:
    """Every part in the zip must be declared in [Content_Types].xml.

    This is the exact defect in the distributed form; the gate keeps it from
    being reintroduced into the deliverable.
    """
    with zipfile.ZipFile(path) as z:
        ct = z.read("[Content_Types].xml").decode("utf-8")
        names = z.namelist()
    declared_ext = set(re.findall(r'<Default Extension="([^"]+)"', ct))
    declared_part = set(re.findall(r'<Override PartName="/([^"]+)"', ct))
    for name in names:
        if name == "[Content_Types].xml" or name in declared_part:
            continue
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext not in declared_ext:
            raise SystemExit(f"undeclared part in package: {name}")


PAGE_MIN, PAGE_MAX = 22, 23


def main() -> int:
    tex = TEX.read_text(encoding="utf-8")
    refs = read_aux(AUX)
    macros = load_macros()

    print("rendering TikZ diagrams from the report source")
    render_tikz(tex)

    blocks = parse_tex(tex, macros)
    kinds = {}
    for b in blocks:
        kinds[b.kind] = kinds.get(b.kind, 0) + 1
    print("blocks:", ", ".join(f"{k}={v}" for k, v in sorted(kinds.items())))

    # Numbering gate: the counters this script derives must agree with the .aux.
    eq_last = max((int(n) for b in blocks if b.kind == "equation"
                   for n in b.numbers if n.isdigit()), default=0)
    if refs.get("eq:bus9power") != str(eq_last):
        raise SystemExit(f"equation numbering disagrees with .aux: "
                         f"derived {eq_last}, aux {refs.get('eq:bus9power')}")
    n_tab = sum(1 for b in blocks if b.kind == "table")
    n_fig = sum(1 for b in blocks if b.kind == "figure")
    if refs.get("tab:runsummary") != str(n_tab) or refs.get("fig:bus9") != str(n_fig):
        raise SystemExit("table/figure numbering disagrees with .aux")
    print(f"numbering gate ok: {n_fig} figures, {n_tab} tables, {eq_last} equations")

    # Page-count search.  Body type stays at the form's 16 pt; only the spacing
    # around content and the raster scale move, down the fixed ladder above.
    # No sentence, table row, figure or heading is removed to reach the target.
    # The walk starts at LAYOUT_START and steps toward whichever end the
    # measurement points to, so the usual case costs one Word round trip.
    trials: list[tuple[int, float, int]] = []
    chosen: Layout | None = None
    math: MathConverter | None = None
    pages = omaths = 0
    idx = LAYOUT_START
    seen: set[int] = set()
    while 0 <= idx < len(LAYOUT_LADDER) and idx not in seen:
        seen.add(idx)
        layout = LAYOUT_LADDER[idx]
        math = MathConverter()
        Builder(layout, math, macros, refs).build(blocks)
        pages, omaths = word_measure(OUT_DOCX)
        trials.append((idx, layout.lead_pt, pages))
        print(f"  step {idx}: lead {layout.lead_pt:.1f} pt, heads x{layout.head_space:.2f}, "
              f"figures x{layout.fig_scale:.2f} -> {pages} pages")
        if PAGE_MIN <= pages <= PAGE_MAX:
            chosen = layout
            break
        idx += 1 if pages > PAGE_MAX else -1
    if chosen is None or not PAGE_MIN <= pages <= PAGE_MAX:
        raise SystemExit(f"page count {pages} outside {PAGE_MIN}-{PAGE_MAX}: {trials}")

    pages, omaths = word_measure(OUT_DOCX, OUT_PDF)
    check_package(OUT_DOCX)

    expected = math.count
    if omaths < expected * 0.98:
        raise SystemExit(f"Word counted {omaths} equations, script emitted {expected}")

    prov = WORDFIG / "provenance.txt"
    prov.write_text(
        "Word build of the Thai project-summary report\n"
        "=============================================\n\n"
        f"source .tex   : {TEX.relative_to(ROOT).as_posix()}\n"
        f"numbering from: {AUX.relative_to(ROOT).as_posix()}\n"
        f"base form     : {FORM.relative_to(ROOT).as_posix()}\n"
        f"output        : {OUT_DOCX.relative_to(ROOT).as_posix()}\n\n"
        f"body font     : {THAI_FONT} {BODY_PT:.1f} pt (the form's own size)\n"
        f"line spacing  : {chosen.lead_pt:.1f} pt exact\n"
        f"heading space : x{chosen.head_space:.2f} of nominal\n"
        f"figure scale  : x{chosen.fig_scale:.2f} of the column-fit width\n"
        f"caption/table : {BODY_PT + chosen.small_delta:.1f} pt\n"
        f"bibliography  : {chosen.bib_columns} column(s)\n"
        f"page count    : {pages} (Word ComputeStatistics, target "
        f"{PAGE_MIN}-{PAGE_MAX})\n"
        f"figures       : {n_fig}\n"
        f"tables        : {n_tab}\n"
        f"equations     : {expected} emitted, {omaths} counted by Word\n"
        f"last eq number: {eq_last}\n\n"
        "layout ladder walked (step, line spacing pt, pages):\n"
        + "".join(f"  {a}  {b:.1f}  {c}\n" for a, b, c in trials)
        + "\nBody type size, wording, headings, tables and figures are fixed.\n"
          "Only the space around content and the raster scale were adjusted to\n"
          "reach the page target; nothing was cut.\n"
          "The two TikZ diagrams are rendered from the report source with\n"
          "xelatex, so their Thai lettering and macro values match the PDF.\n"
          "Raster figures are the PNGs the MATLAB generator produced, scaled to\n"
          "the form's text column and otherwise unmodified.\n"
          "No numeric value was recomputed, rounded or reformatted here.\n",
        encoding="utf-8")

    print(f"\nwrote {OUT_DOCX.relative_to(ROOT).as_posix()}")
    print(f"wrote {OUT_PDF.relative_to(ROOT).as_posix()}")
    print(f"wrote {prov.relative_to(ROOT).as_posix()}")
    print(f"pages {pages}, equations {expected} (Word: {omaths}), "
          f"body {BODY_PT:.1f} pt, lead {chosen.lead_pt:.1f} pt")
    return 0


if __name__ == "__main__":
    sys.exit(main())
